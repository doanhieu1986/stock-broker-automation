"""
skills/eod_report/scripts/generate_eod_docx.py
Bước 2 ca EOD: Tạo báo cáo cuối ngày dạng Word.
"""

import sys
from datetime import datetime
from pathlib  import Path
sys.path.insert(0, str(Path(__file__).parents[4]))
from docx import Document
from docx.shared import Pt, RGBColor

from utils.data_loader import load_cache, load_alerts, output_dir, today
from utils.logger      import get_logger

log = get_logger(__name__)

TEMPLATE  = Path("templates/eod_report.docx")
DISCLAIMER = "Thông tin chỉ mang tính tham khảo, không phải khuyến nghị đầu tư."


def run() -> Path:
    """Tạo file Word báo cáo EOD từ dữ liệu đóng cửa đã thu thập."""
    log.info("Tạo báo cáo EOD Word...")
    date_str = today()
    eod      = load_cache("eod_raw", date_str) or {}

    if not eod:
        log.error("Không có dữ liệu EOD — chạy collect_eod_data.py trước")
        return None

    doc = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    if TEMPLATE.exists():
        for p in doc.paragraphs[1:]:
            p.clear()

    date_show = datetime.now().strftime("%d/%m/%Y")

    # ── Section 1: Dashboard tổng quan ────────────────────────────────────
    doc.add_heading(f"BÁO CÁO PHIÊN GIAO DỊCH — {date_show}", 1)

    idx = eod.get("indices", {})
    _add_index_table(doc, idx)

    foreign = eod.get("foreign_flow", {})
    if foreign:
        net = foreign.get("net_bil", 0)
        sign = "Mua ròng" if net >= 0 else "Bán ròng"
        doc.add_paragraph(
            f"Khối ngoại: {sign} {abs(net):,.1f} tỷ VNĐ "
            f"(Mua: {foreign.get('buy_bil',0):,.1f} / Bán: {foreign.get('sell_bil',0):,.1f})"
        )
    doc.add_page_break()

    # ── Section 2: Phân tích thị trường ──────────────────────────────────
    doc.add_heading("Phân tích thị trường hôm nay", 2)
    doc.add_paragraph(_market_narrative(eod))

    doc.add_heading("Top movers HOSE", 3)
    _add_movers_table(doc, eod.get("top_gainers", []), eod.get("top_losers", []))
    doc.add_page_break()

    # ── Section 3: Watchlist hiệu suất ────────────────────────────────────
    doc.add_heading("Watchlist — Hiệu suất cả ngày", 2)
    wl = eod.get("watchlist_eod", [])
    if wl:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Mã", "Đóng cửa", "Thay đổi %", "KL (K cp)"]):
            table.rows[0].cells[i].text = h
        for item in sorted(wl, key=lambda x: x["change_pct"], reverse=True):
            row = table.add_row().cells
            row[0].text = item["ticker"]
            row[1].text = f"{item['close']:,.1f}"
            row[2].text = f"{item['change_pct']:+.2f}%"
            row[3].text = f"{item['volume']:,}"
    doc.add_page_break()

    # ── Section 4: Khuyến nghị (broker điền) ──────────────────────────────
    doc.add_heading("Khuyến nghị (Broker điền tay)", 2)
    doc.add_paragraph("⚠ Phần này do Broker điền sau khi review — Claude để trống.")
    reco_table = doc.add_table(rows=4, cols=4)
    reco_table.style = "Table Grid"
    for i, h in enumerate(["Mã CK", "Hành động", "Lý do", "Giá tham chiếu"]):
        reco_table.rows[0].cells[i].text = h

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph(
        f"⚠ {DISCLAIMER}\n"
        f"Tạo lúc: {datetime.now().strftime('%H:%M:%S %d/%m/%Y')} | Chờ Broker duyệt"
    )
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    out_path = output_dir(date_str) / f"eod_report_{date_str}.docx"
    doc.save(str(out_path))
    log.info(f"Đã lưu: {out_path}")
    return out_path


def _add_index_table(doc: Document, idx: dict) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Chỉ số", "Đóng cửa", "Thay đổi pt", "Thay đổi %"]):
        table.rows[0].cells[i].text = h
    for key, label in [("vn_index","VN-Index"),("hnx_index","HNX-Index"),
                        ("vn30","VN30"),("upcom_index","UPCOM")]:
        d = idx.get(key, {})
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{d.get('value', 'N/A')}"
        row[2].text = f"{d.get('change_pt', 'N/A'):+.2f}" if d.get('change_pt') is not None else "N/A"
        row[3].text = f"{d.get('change_pct', 'N/A'):+.2f}%" if d.get('change_pct') is not None else "N/A"
    doc.add_paragraph("")


def _add_movers_table(doc: Document, gainers: list, losers: list) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Top tăng", "% Tăng", "Top giảm", "% Giảm"]):
        table.rows[0].cells[i].text = h
    for i in range(max(len(gainers), len(losers))):
        g = gainers[i] if i < len(gainers) else {}
        l = losers[i]  if i < len(losers)  else {}
        row = table.add_row().cells
        row[0].text = g.get("ticker", "")
        row[1].text = f"{g.get('change_pct', 0):+.2f}%" if g else ""
        row[2].text = l.get("ticker", "")
        row[3].text = f"{l.get('change_pct', 0):+.2f}%" if l else ""


def _market_narrative(eod: dict) -> str:
    """Tạo nhận định 3-5 câu từ dữ liệu EOD."""
    idx = eod.get("indices", {})
    vn  = idx.get("vn_index", {})
    chg = vn.get("change_pct", 0)
    vol = vn.get("volume_bil", 0)
    direction = "tăng điểm" if chg >= 0 else "giảm điểm"

    foreign = eod.get("foreign_flow", {})
    net = foreign.get("net_bil", 0)
    foreign_text = (
        f"Khối ngoại {'mua' if net >= 0 else 'bán'} ròng {abs(net):,.1f} tỷ VNĐ."
    )

    gainers = eod.get("top_gainers", [])
    top_g   = gainers[0]["ticker"] if gainers else "N/A"
    top_chg = gainers[0].get("change_pct", 0) if gainers else 0

    return (
        f"VN-Index {direction} {chg:+.2f}%, thanh khoản đạt {vol:,.0f} tỷ VNĐ. "
        f"Cổ phiếu dẫn dắt mạnh nhất: {top_g} (+{top_chg:.2f}%). "
        f"{foreign_text}"
    )
