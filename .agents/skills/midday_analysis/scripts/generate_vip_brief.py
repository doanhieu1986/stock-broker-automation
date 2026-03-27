"""
skills/midday_analysis/scripts/generate_vip_brief.py
Bước 3 ca trưa: Tạo báo cáo Word cá nhân hóa cho KH VIP.
"""

from datetime import datetime
from pathlib  import Path
from docx     import Document
from docx.shared import Pt, RGBColor
from utils.api_client  import fetch_cafef_quote
from utils.data_loader import load_customers, load_cache, output_dir, today
from utils.logger      import get_logger

log = get_logger(__name__)

DISCLAIMER = ("Báo cáo chỉ nhằm mục đích thông tin. "
              "Không phải khuyến nghị đầu tư. "
              "Nhà đầu tư tự chịu trách nhiệm về quyết định của mình.")


def run() -> list[Path]:
    log.info("Tạo báo cáo VIP...")
    date_str    = today()
    vip_list    = load_customers(vip_only=True)
    morning_sum = load_cache("morning_session", date_str) or {}
    out_paths   = []

    for customer in vip_list:
        path = _generate(customer, date_str, morning_sum)
        if path:
            out_paths.append(path)
            log.info(f"  {customer['name']} → {path.name}")

    log.info(f"  {len(out_paths)} báo cáo VIP — chờ broker review")
    return out_paths


def _generate(customer: dict, date_str: str, morning_sum: dict) -> Path | None:
    name     = customer.get("name", "Quý khách")
    holdings = customer.get("holdings", [])
    doc      = Document()

    doc.add_heading(f"CẬP NHẬT THỊ TRƯỜNG — {datetime.now().strftime('%d/%m/%Y')}", 1)
    doc.add_paragraph(f"Kính gửi: {name}")

    doc.add_heading("1. Danh mục hôm nay", 2)
    total_pnl = 0.0
    if holdings:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Mã", "Giá vốn", "Giá hiện tại", "P&L hôm nay"]):
            table.rows[0].cells[i].text = h
        for h in holdings:
            ticker = h.get("ticker", "")
            cost   = h.get("cost", 0)
            shares = h.get("shares", 0)
            quote  = fetch_cafef_quote(ticker) or {}
            cur_p  = quote.get("price", cost)
            chg    = quote.get("change_pct", 0)
            pnl    = (cur_p - cost) * shares
            total_pnl += pnl
            row = table.add_row().cells
            row[0].text = ticker
            row[1].text = f"{cost:,.1f}"
            row[2].text = f"{cur_p:,.1f}  ({chg:+.2f}%)"
            row[3].text = f"{pnl:+,.0f} VNĐ"
    doc.add_paragraph(f"\nTổng P&L hôm nay: {total_pnl:+,.0f} VNĐ")

    doc.add_heading("2. Thị trường hôm nay", 2)
    commentary = morning_sum.get("commentary", "")
    if commentary:
        doc.add_paragraph(commentary)

    doc.add_heading("3. Broker sẽ liên hệ", 2)
    doc.add_paragraph("[Broker điền: ___h___ qua □ Điện thoại  □ Zalo]")

    p = doc.add_paragraph(DISCLAIMER)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    out_dir  = output_dir(date_str) / "private"
    out_dir.mkdir(exist_ok=True)
    safe     = customer.get("id", name.replace(" ", "_"))
    out_path = out_dir / f"vip_{safe}_{date_str}.docx"
    doc.save(str(out_path))
    return out_path
