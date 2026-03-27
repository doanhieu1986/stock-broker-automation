"""
skills/morning-prep/scripts/generate_morning_pdf.py
Bước 5 ca sáng (8h30–8h45): Tổng hợp toàn bộ → xuất PDF + tóm tắt terminal.
Output: outputs/{date}/morning_{date}.pdf
        outputs/{date}/morning_{date}_summary.txt
"""

import sys
import json
from datetime import datetime, date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from docx2pdf import convert
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

sys.path.insert(0, str(Path(__file__).parents[3]))
from utils.logger import section, Timer, log_output, print_box
from utils.api_helpers import (
    cache_read, output_dir, load_watchlist, fmt_pct,
)

DISCLAIMER = "Dữ liệu chỉ mang tính tham khảo. Không phải khuyến nghị đầu tư."
TEMPLATE   = Path("templates/morning_report.docx")


# ── Kiểm tra chất lượng dữ liệu ─────────────────────────────────────────────

def _quality_check(global_d: dict, asia_d: dict, fund_d: dict) -> list[str]:
    """Trả về danh sách cảnh báo. [] = pass."""
    warns = []
    indices = global_d.get("indices", {}) if global_d else {}
    ok_count = sum(1 for v in indices.values() if v.get("price") is not None)
    if ok_count < 3:
        warns.append(f"Chỉ có {ok_count}/3 chỉ số Mỹ — kiểm tra Yahoo Finance")

    markets = asia_d.get("markets", {}) if asia_d else {}
    ok_asia = sum(1 for v in markets.values() if v.get("price") is not None)
    if ok_asia < 3:
        warns.append(f"Chỉ có {ok_asia}/5 thị trường châu Á")

    if not fund_d or not fund_d.get("data"):
        warns.append("Thiếu dữ liệu phân tích cơ bản")

    miss_pct = (5 - ok_count + 5 - ok_asia) / 10
    if miss_pct > 0.3:
        warns.append(f"⚠ CHẤT LƯỢNG THẤP: {miss_pct*100:.0f}% dữ liệu bị thiếu")
    return warns


# ── Tạo file Word ─────────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_table_row(table, cells: list, bold_first: bool = False):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = str(val)
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True


def _color_change(cell, pct: float | None):
    if pct is None:
        return
    run = cell.paragraphs[0].runs
    if not run:
        return
    run[0].font.color.rgb = RGBColor(0, 153, 51) if pct >= 0 else RGBColor(204, 0, 0)


def build_docx(
    global_d: dict,
    asia_d:   dict,
    reg_d:    dict,
    fund_d:   dict,
    ds:       str,
) -> Path:
    """Tạo file Word và trả về Path."""
    doc = Document(TEMPLATE) if TEMPLATE.exists() else Document()

    # ── Trang 1: Tổng quan ────────────────────────────────────────────────────
    _heading(doc, f"BÁO CÁO SÁNG — {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("")

    # Bảng thị trường Mỹ
    _heading(doc, "Thị trường Mỹ & Hàng hóa", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Chỉ số", "Giá", "% Thay đổi"

    if global_d:
        for info in global_d.get("indices", {}).values():
            row = tbl.add_row().cells
            row[0].text = info["label"]
            row[1].text = f"{info['price']:,.2f}" if info["price"] else "N/A"
            chg = info.get("change_pct")
            row[2].text = fmt_pct(chg) if chg is not None else "N/A"
            _color_change(row[2], chg)
        for info in global_d.get("commodities", {}).values():
            row = tbl.add_row().cells
            row[0].text = info["label"]
            row[1].text = f"${info['price']:,.1f}" if info["price"] else "N/A"
            chg = info.get("change_pct")
            row[2].text = fmt_pct(chg) if chg is not None else "N/A"
            _color_change(row[2], chg)
        usd = global_d.get("fx", {}).get("usdvnd", {}).get("rate")
        if usd:
            row = tbl.add_row().cells
            row[0].text, row[1].text, row[2].text = "USD/VND", f"{usd:,.0f}", "—"

    doc.add_paragraph("")

    # Bảng châu Á
    _heading(doc, "Thị trường châu Á", level=2)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    h2 = tbl2.rows[0].cells
    h2[0].text, h2[1].text, h2[2].text = "Thị trường", "Điểm", "% Thay đổi"

    if asia_d:
        for info in asia_d.get("markets", {}).values():
            row = tbl2.add_row().cells
            flag = info.get("flag") or ""
            row[0].text = f"{flag} {info['label']}"
            row[1].text = f"{info['price']:,.2f}" if info["price"] else "N/A"
            chg = info.get("change_pct")
            row[2].text = fmt_pct(chg) if chg is not None else "(đóng cửa hôm qua)"
            _color_change(row[2], chg)

    doc.add_page_break()

    # ── Trang 2: Tin tức & sự kiện ────────────────────────────────────────────
    _heading(doc, "Tin tức & Sự kiện")

    if reg_d:
        ubck = reg_d.get("ubcknn", [])
        if ubck:
            _heading(doc, "Thông báo UBCKNN (24h qua)", level=2)
            for item in ubck[:5]:
                flag = "⚠ " if item.get("priority") else "• "
                p = doc.add_paragraph(f"{flag}{item['title']}")
                p.runs[0].font.size = Pt(10)
        else:
            doc.add_paragraph("Không có thông báo UBCKNN mới.")

        events = reg_d.get("events", [])
        if events:
            _heading(doc, "Sự kiện doanh nghiệp (3 ngày tới)", level=2)
            tbl3 = doc.add_table(rows=1, cols=4)
            tbl3.style = "Table Grid"
            h3 = tbl3.rows[0].cells
            h3[0].text, h3[1].text, h3[2].text, h3[3].text = "Ngày", "Mã", "Loại", "Nội dung"
            for ev in events[:8]:
                wl = "★ " if ev.get("in_watchlist") else ""
                row = tbl3.add_row().cells
                row[0].text = ev["date"]
                row[1].text = f"{wl}{ev['ticker']}"
                row[2].text = ev["event_label"]
                row[3].text = ev["description"][:40]

    doc.add_page_break()

    # ── Trang 3: Phân tích cơ bản ─────────────────────────────────────────────
    _heading(doc, "Phân tích cơ bản")

    if fund_d and fund_d.get("data"):
        pair   = fund_d["pair"]
        sector = fund_d["sector"]
        a_dat  = fund_d["data"].get(pair[0], {})
        b_dat  = fund_d["data"].get(pair[1], {})
        _heading(doc, f"{pair[0]} vs {pair[1]} — Ngành: {sector}", level=2)

        from skills.morning_prep.scripts.fundamental_compare import METRICS
        tbl4 = doc.add_table(rows=1, cols=3)
        tbl4.style = "Table Grid"
        h4 = tbl4.rows[0].cells
        h4[0].text, h4[1].text, h4[2].text = "Chỉ số", pair[0], pair[1]

        for field, label, fmt in METRICS:
            av, bv = a_dat.get(field), b_dat.get(field)
            row = tbl4.add_row().cells
            row[0].text = label
            try:
                row[1].text = fmt.format(av) if av is not None else "N/A"
                row[2].text = fmt.format(bv) if bv is not None else "N/A"
            except Exception:
                row[1].text = str(av or "N/A")
                row[2].text = str(bv or "N/A")

        doc.add_paragraph("")
        p = doc.add_paragraph(fund_d.get("narrative", ""))
        p.runs[0].italic = True if p.runs else None

    # Footer disclaimer
    doc.add_page_break()
    footer_p = doc.add_paragraph(
        f"\n{DISCLAIMER}\nTạo lúc {datetime.now().strftime('%H:%M %d/%m/%Y')}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_p.runs:
        footer_p.runs[0].font.size = Pt(9)
        footer_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Lưu
    out = output_dir(ds)
    docx_path = out / f"morning_{ds}.docx"
    doc.save(docx_path)
    return docx_path


# ── Terminal summary ──────────────────────────────────────────────────────────

def _terminal_summary(global_d, asia_d, reg_d, ds: str) -> str:
    now  = datetime.now().strftime("%H:%M")
    rows = []

    # Global
    if global_d:
        indices = global_d.get("indices", {})
        dow  = indices.get("dow", {})
        sp   = indices.get("sp500", {})
        vix  = indices.get("vix", {})
        rows.append("🌍 QUỐC TẾ")
        if dow.get("price"):
            rows.append(f"   Dow {dow['price']:,.0f} ({fmt_pct(dow['change_pct'])})  "
                        f"S&P {sp.get('price', 'N/A'):,.0f}")
        if vix.get("price"):
            mood = "thấp" if vix["price"] < 15 else ("TB" if vix["price"] < 25 else "cao")
        rows.append(f"   VIX {vix.get('price','N/A'):.1f} ({mood})" if vix.get("price") else "")

    # Asia
    if asia_d:
        rows.append("🌏 CHÂU Á")
        for info in asia_d.get("markets", {}).values():
            chg  = f"({fmt_pct(info['change_pct'])})" if info.get("change_pct") is not None else "(—)"
            flag = info.get("flag") or "  "
            rows.append(f"   {flag} {info['label']:<22} {chg}")

    # Alerts
    if reg_d:
        alerts = reg_d.get("priority_alerts", [])
        if alerts:
            rows.append("⚠️  ĐÁNG CHÚ Ý")
            for a in alerts[:3]:
                t = a.get("title") or f"{a.get('ticker','')} — {a.get('event_label','')}"
                rows.append(f"   • {t[:55]}")

    rows.append(f"📁 outputs/{ds}/morning_{ds}.pdf")

    txt = "\n".join(r for r in rows if r)
    out = output_dir(ds)
    summary_path = out / f"morning_{ds}_summary.txt"
    summary_path.write_text(txt, encoding="utf-8")
    return txt


# ── Main ─────────────────────────────────────────────────────────────────────

def run() -> Path:
    ds  = date.today().strftime("%Y%m%d")
    section("Bước 5 — Tổng hợp & Xuất báo cáo")

    # Load từ cache
    global_d = cache_read("global")
    asia_d   = cache_read("asia")
    reg_d    = cache_read("regulatory")
    fund_d   = cache_read("fundamental")

    # Quality check
    warns = _quality_check(global_d, asia_d, fund_d)
    if warns:
        print("\n  ⚠ CẢNH BÁO CHẤT LƯỢNG DỮ LIỆU:")
        for w in warns:
            print(f"    • {w}")

    with Timer("Tạo file Word"):
        docx_path = build_docx(global_d, asia_d, reg_d, fund_d, ds)
    log_output(str(docx_path), docx_path.stat().st_size / 1024)

    # Convert sang PDF nếu có docx2pdf
    pdf_path = docx_path.with_suffix(".pdf")
    if HAS_PDF:
        with Timer("Convert → PDF"):
            convert(str(docx_path), str(pdf_path))
        log_output(str(pdf_path), pdf_path.stat().st_size / 1024)
        print(f"\n  ✓ PDF: {pdf_path}")
    else:
        print(f"\n  ✓ DOCX: {docx_path}  (cài docx2pdf để convert sang PDF)")
        pdf_path = docx_path

    # Terminal summary
    summary = _terminal_summary(global_d, asia_d, reg_d, ds)
    print_box(
        f"BÁO CÁO SÁNG — {datetime.now().strftime('%d/%m/%Y')} — Tạo lúc {datetime.now().strftime('%H:%M')}",
        summary.split("\n"),
    )
    return pdf_path


if __name__ == "__main__":
    run()
